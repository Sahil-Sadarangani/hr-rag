from dataclasses import dataclass
from pathlib import Path
from typing import List

from pypdf import PdfReader
from docx import Document as DocxDocument


@dataclass
class LoadedDocument:
    source: str
    text: str


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for page_number, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(f"[Page {page_number}]\n{page_text}")
    return "\n\n".join(pages)


def _read_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def load_single_document(path: Path) -> LoadedDocument:
    suffix = path.suffix.lower()

    if suffix in {".txt", ".md"}:
        text = _read_txt(path)
    elif suffix == ".pdf":
        text = _read_pdf(path)
    elif suffix == ".docx":
        text = _read_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {path.name}")

    return LoadedDocument(source=path.name, text=text)


def load_documents(data_dir: Path) -> List[LoadedDocument]:
    if not data_dir.exists():
        raise FileNotFoundError(f"Data directory not found: {data_dir}")

    files = [p for p in data_dir.iterdir() if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS]
    if not files:
        raise FileNotFoundError(
            f"No supported documents found in {data_dir}. "
            f"Supported types: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    documents: List[LoadedDocument] = []
    for file_path in sorted(files):
        loaded = load_single_document(file_path)
        if loaded.text.strip():
            documents.append(loaded)

    return documents
